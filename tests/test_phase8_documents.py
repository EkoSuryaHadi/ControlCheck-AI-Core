from __future__ import annotations

from pathlib import Path
from docx import Document

ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"


def _docx_text(path: Path) -> str:
    document = Document(path)
    blocks = [paragraph.text for paragraph in document.paragraphs]
    blocks.extend(cell.text for table in document.tables for row in table.rows for cell in row.cells)
    return "\n".join(blocks)


def test_prd_v09_records_web_dashboard_delivery() -> None:
    text = _docx_text(DOCS / "ControlCheck_AI_PRD_v0.9.docx")
    assert "Product Requirements Document v0.9" in text
    assert "Phase 8 Web Dashboard & AI Copilot Alignment" in text
    assert "Change Log v0.9" in text
    assert "Interactive Single-Page Application" in text
