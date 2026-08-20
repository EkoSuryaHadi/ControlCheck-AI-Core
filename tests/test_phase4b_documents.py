from __future__ import annotations

from pathlib import Path
from docx import Document

ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"


def _docx_text(path: Path) -> str:
    document = Document(path)
    blocks = [paragraph.text for paragraph in document.paragraphs]
    blocks.extend(cell.text for table in document.tables for row in table.rows for cell in row.cells)
    return "\n".join(blocks)


def test_phase4b_sql_reference() -> None:
    text = (DOCS / "003_controlcheck_canonical_schema_v0.3.sql").read_text(encoding="utf-8")
    assert "Phase 4B Canonical Facts" in text
    for table in (
        "raw_rows",
        "import_batches",
        "import_column_mappings",
        "wbs_nodes",
        "budget_records",
        "cost_records",
        "commitment_records",
        "schedule_activities",
        "progress_records",
    ):
        assert f"CREATE TABLE {table}" in text


def test_erd_v03_documents_phase4b_canonical_facts() -> None:
    text = _docx_text(DOCS / "ControlCheck_AI_ERD_Database_Spec_v0.3.docx")
    assert "ERD & Database Specification v0.3" in text
    assert "Phase 4B Canonical Facts Alignment" in text
    for table in ("raw_rows", "import_batches", "import_column_mappings", "wbs_nodes", "budget_records", "cost_records"):
        assert table in text
    assert "raw_row_id" in text


def test_prd_v04_records_phase4b_delivery() -> None:
    text = _docx_text(DOCS / "ControlCheck_AI_PRD_v0.4.docx")
    assert "Product Requirements Document v0.4" in text
    assert "Phase 4B Product Alignment" in text
    assert "Change Log v0.4" in text
