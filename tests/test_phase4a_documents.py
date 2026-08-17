from __future__ import annotations

import json
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"


def _docx_text(path: Path) -> str:
    document = Document(path)
    blocks = [paragraph.text for paragraph in document.paragraphs]
    blocks.extend(cell.text for table in document.tables for row in table.rows for cell in row.cells)
    return "\n".join(blocks)


def test_phase4a_sql_is_an_alembic_aligned_readable_reference() -> None:
    text = (DOCS / "002_controlcheck_persistence_schema_v0.2.sql").read_text(encoding="utf-8")
    assert "Alembic" in text
    assert "executable authority" in text
    for table in (
        "organizations",
        "projects",
        "source_files",
        "dataset_snapshots",
        "rule_catalogue_versions",
        "analysis_runs",
        "findings",
        "finding_evidence",
        "approved_exceptions",
        "audit_logs",
    ):
        assert f"CREATE TABLE {table}" in text


def test_erd_v02_documents_phase4a_persistence_contract() -> None:
    text = _docx_text(DOCS / "ControlCheck_AI_ERD_Database_Spec_v0.2.docx")
    assert "ERD & Database Specification v0.2" in text
    for table in ("dataset_snapshots", "rule_catalogue_versions", "analysis_runs", "approved_exceptions"):
        assert table in text
    assert "entity_id" in text and "text" in text.lower()
    assert "Phase 4A" in text


def test_rule_catalogue_v02_matches_runtime_catalogue() -> None:
    catalogue = json.loads((ROOT / "data" / "controlcheck_rule_catalogue_v0.2.json").read_text(encoding="utf-8"))
    text = _docx_text(DOCS / "ControlCheck_AI_Control_Rule_Catalogue_v0.2.docx")
    assert "Control Rule Catalogue v0.2" in text
    assert len(catalogue["rules"]) == 20
    for rule in catalogue["rules"]:
        assert rule["code"] in text
    assert "vendor + WBS" in text
    assert "25% of WBS budget" in text
    assert "3% of project budget" in text
    assert "1% of project budget" in text


def test_prd_v03_records_phase_sequencing_and_auth_deferral() -> None:
    text = _docx_text(DOCS / "ControlCheck_AI_PRD_v0.3.docx")
    assert "Product Requirements Document v0.3" in text
    assert "Phase 4A" in text
    assert "Phase 4B" in text
    assert "Authentication and complete RBAC are deferred" in text
    assert "Change Log" in text

