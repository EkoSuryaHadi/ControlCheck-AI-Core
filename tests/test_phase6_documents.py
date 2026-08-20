from __future__ import annotations

from pathlib import Path
from docx import Document

ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"


def _docx_text(path: Path) -> str:
    document = Document(path)
    blocks = [paragraph.text for paragraph in document.paragraphs]
    blocks.extend(cell.text for table in document.tables for row in table.rows for cell in row.cells)
    return "\n".join(blocks)


def test_prd_v07_records_ai_layer_delivery() -> None:
    text = _docx_text(DOCS / "ControlCheck_AI_PRD_v0.7.docx")
    assert "Product Requirements Document v0.7" in text
    assert "Phase 6 Product Alignment" in text
    assert "Change Log v0.7" in text
    assert "AI Intelligence Layer" in text
    assert "Deterministic engine calculates, AI explains" in text
