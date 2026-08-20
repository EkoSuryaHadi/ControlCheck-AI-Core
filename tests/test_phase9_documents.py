from __future__ import annotations

from pathlib import Path
from docx import Document

ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"


def _docx_text(path: Path) -> str:
    document = Document(path)
    blocks = [paragraph.text for paragraph in document.paragraphs]
    blocks.extend(cell.text for table in document.tables for row in table.rows for cell in row.cells)
    return "\n".join(blocks)


def test_prd_v10_records_universal_flexible_ingestion() -> None:
    text = _docx_text(DOCS / "ControlCheck_AI_PRD_v1.0.docx")
    assert "Product Requirements Document v1.0" in text
    assert "Phase 9 Smart Flexible Ingestion & Universal Auto-Mapper Alignment" in text
    assert "Change Log v1.0" in text
    assert "Fuzzy sheet name detection" in text
