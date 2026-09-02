from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]


def test_prd_v13_records_ai_insight_beta_scope() -> None:
    document = Document(ROOT / "docs" / "ControlCheck_AI_PRD_v1.3.docx")
    text = "\n".join(
        [paragraph.text for paragraph in document.paragraphs]
        + [cell.text for table in document.tables for row in table.rows for cell in row.cells]
    )
    assert "Product Requirements Document v1.3" in text
    assert "AI Insight v1" in text
    assert "workbook raw content is not sent" in text.lower()