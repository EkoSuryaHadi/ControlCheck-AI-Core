import hashlib
from pathlib import Path

from docx import Document


REQUIRED_PRODUCT_CHANGES = [
    "Rule and Threshold Governance",
    "Approved Exceptions",
    "Ground-Truth Governance",
    "Validation Metrics",
    "Artifact Version Compatibility",
    "Change Log v0.2",
]

PRD_V01_SHA256 = "1705961a435502289c643251816ea343209f62c2e4460642908e01d425f51c91"


def _all_text(path: Path) -> str:
    document = Document(path)
    blocks = [paragraph.text for paragraph in document.paragraphs]
    blocks.extend(cell.text for table in document.tables for row in table.rows for cell in row.cells)
    return "\n".join(blocks)


def _sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def test_prd_v02_contains_accepted_product_changes(project_root: Path):
    prd_v02 = project_root / "docs" / "ControlCheck_AI_PRD_v0.2.docx"
    text = _all_text(prd_v02)

    assert all(section in text for section in REQUIRED_PRODUCT_CHANGES)
    assert "100% on controlled validation fixtures" in text
    assert "not a customer-accuracy claim" in text
    assert "CST-004" in text and "WBS|VENDOR" in text
    assert "CST-005" in text and "25%" in text and "3%" in text
    assert "PRG-003" in text and "1%" in text
    assert "incompatible_artifact_versions" in text


def test_prd_version_metadata_and_change_log_are_updated(project_root: Path):
    prd_v02 = project_root / "docs" / "ControlCheck_AI_PRD_v0.2.docx"
    document = Document(prd_v02)
    text = _all_text(prd_v02)

    assert "Version 0.2 | MVP Blueprint | 17 August 2026" in text
    assert document.core_properties.version == "0.2"
    change_log = document.tables[-1]
    rows = [[cell.text for cell in row.cells] for row in change_log.rows]
    assert any(row[0] == "0.2" and "Validation Alignment" in row[2] for row in rows)


def test_prd_v01_is_preserved_byte_for_byte(project_root: Path):
    preserved = project_root / "docs" / "ControlCheck_AI_PRD_v0.1.docx"

    assert _sha256(preserved) == PRD_V01_SHA256


def test_prd_v02_has_no_internal_placeholders(project_root: Path):
    text = _all_text(project_root / "docs" / "ControlCheck_AI_PRD_v0.2.docx")
    forbidden = [":codex-file-citation", "TODO", "TBD", "{{", "}}"]
    assert not any(token in text for token in forbidden)

