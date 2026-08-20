from __future__ import annotations

from pathlib import Path
from docx import Document

ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"


def _docx_text(path: Path) -> str:
    document = Document(path)
    blocks = [paragraph.text for paragraph in document.paragraphs]
    blocks.extend(cell.text for table in document.tables for row in table.rows for cell in row.cells)
    return "\n".join(blocks)


def test_prd_v05_records_health_scoring_and_api_hardening() -> None:
    text = _docx_text(DOCS / "ControlCheck_AI_PRD_v0.5.docx")
    assert "Product Requirements Document v0.5" in text
    assert "Phase 5B & 5C Product Alignment" in text
    assert "Change Log v0.5" in text
    assert "Deterministic health scoring engine" in text
