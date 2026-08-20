from __future__ import annotations

from pathlib import Path
from docx import Document

ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"


def _docx_text(path: Path) -> str:
    document = Document(path)
    blocks = [paragraph.text for paragraph in document.paragraphs]
    blocks.extend(cell.text for table in document.tables for row in table.rows for cell in row.cells)
    return "\n".join(blocks)


def test_prd_v06_records_auth_and_rbac_delivery() -> None:
    text = _docx_text(DOCS / "ControlCheck_AI_PRD_v0.6.docx")
    assert "Product Requirements Document v0.6" in text
    assert "Phase 4C Product Alignment" in text
    assert "Change Log v0.6" in text
    assert "JWT authentication" in text
