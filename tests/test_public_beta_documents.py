from __future__ import annotations

from pathlib import Path
import subprocess
import sys
import time

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"


def _docx_text(path: Path) -> str:
    document = Document(path)
    blocks = [paragraph.text for paragraph in document.paragraphs]
    blocks.extend(cell.text for table in document.tables for row in table.rows for cell in row.cells)
    return "\n".join(blocks)


def test_prd_v11_records_public_beta_cloud_deployment_and_usage_validation() -> None:
    text = _docx_text(DOCS / "ControlCheck_AI_PRD_v1.1.docx")

    for expected in (
        "Product Requirements Document v1.1",
        "Version 1.1 | Public Beta Cloud Deployment & Usage Validation | 26 August 2026",
        "Phase 10 Public Beta Cloud Deployment & Usage Validation Alignment",
        "browser → Vercel frontend → Render Free FastAPI → Supabase Free PostgreSQL plus private Cloudflare R2 Standard object storage",
        "register/login → create project → upload workbook → persist workbook in R2 → canonical ingestion and deterministic analysis on Render → persist run/findings/evidence in Supabase → display results in Vercel",
        "hosted register-to-findings flow passes",
        "uploaded files and results persist across Render restart/cold start",
        "no secrets appear in source/logs/docs",
        "registrations, active users, projects, workbook uploads, and completed analysis runs are measurable from persisted records",
        "Render cold start after idle",
        "Supabase Free may pause after low activity and has limited capacity/no managed downloadable backups",
        "R2 Standard free allowance is used and the bucket remains private",
        "Supabase database approaches 400 MB",
        "R2 approaches 8 GB",
        "full authentication/RBAC hardening, payment/subscription, enterprise SSO, and production-scale HA/DR remain deferred",
        "Change Log v1.1",
        "26 Aug 2026",
        "public-beta cloud architecture and usage validation",
    ):
        assert expected in text


def test_public_beta_operations_docs_record_exact_architecture_and_recovery_contract() -> None:
    runbook = (DOCS / "PRODUCTION_RUNBOOK.md").read_text(encoding="utf-8")
    readme = (ROOT / "README.md").read_text(encoding="utf-8")
    architecture = "browser → Vercel frontend → Render Free FastAPI → Supabase Free PostgreSQL plus private Cloudflare R2 Standard object storage"

    for expected in (
        architecture,
        "controlcheck-api",
        "Singapore",
        "/health/ready",
        "CONTROLCHECK_DATABASE_URL",
        "Session Pooler port 5432",
        "controlcheck-beta-workbooks",
        "Standard/private",
        "https://control-check-ai-git-codex-public-8a91d7-ekosuryahadis-projects.vercel.app",
        "controlcheck-api.onrender.com",
        "VITE_API_BASE_URL",
        "Render cold start after idle",
        "Supabase Free may pause after low activity",
        "no managed downloadable backups",
        "private",
        "no secrets appear in source/logs/docs",
        "hosted register-to-findings flow passes",
        "uploaded files and results persist across Render restart/cold start",
        "Create the Supabase Free project",
        "Create the private R2 Standard bucket",
        "Deploy the Render `controlcheck-api` Blueprint",
        "Set Vercel `VITE_API_BASE_URL`",
        "registrations, active users, projects, workbook uploads, and completed analysis runs",
        "region `auto`",
        "Supabase database approaches 400 MB",
        "R2 approaches 8 GB",
        "full authentication/RBAC hardening, payment/subscription, enterprise SSO, and production-scale HA/DR remain deferred",
    ):
        assert expected in runbook

    for expected in (
        architecture,
        "Vercel is frontend-only",
        "Render Free FastAPI",
        "Supabase Free PostgreSQL",
        "Cloudflare R2 Standard",
        "VITE_API_BASE_URL",
        "CONTROLCHECK_DATABASE_URL",
        "controlcheck-beta-workbooks",
        "Render cold start after idle",
        "Supabase Free may pause after low activity",
        "no secrets appear in source/logs/docs",
        "register/login → create project → upload workbook",
        "Provision in that order: Supabase database, private R2 bucket and scoped credentials, Render service plus secrets, then Vercel `VITE_API_BASE_URL`",
        "Registrations, active users, projects, workbook uploads, and completed analysis runs",
        "region `auto`",
        "Supabase database approaches 400 MB",
        "R2 approaches 8 GB",
        "Full authentication/RBAC hardening, payment/subscription, enterprise SSO, and production-scale HA/DR remain deferred",
    ):
        assert expected in readme


def test_public_beta_document_generator_and_plan_are_versioned() -> None:
    assert (ROOT / "tools" / "update_public_beta_documents.py").is_file()
    assert (DOCS / "superpowers" / "plans" / "2026-08-26-public-beta-cloud-deployment.md").is_file()


def test_public_beta_prd_generator_is_archive_byte_deterministic() -> None:
    generator = ROOT / "tools" / "update_public_beta_documents.py"
    target = DOCS / "ControlCheck_AI_PRD_v1.1.docx"

    subprocess.run([sys.executable, str(generator)], cwd=ROOT, check=True)
    first = target.read_bytes()
    time.sleep(2.1)
    subprocess.run([sys.executable, str(generator)], cwd=ROOT, check=True)

    assert target.read_bytes() == first
