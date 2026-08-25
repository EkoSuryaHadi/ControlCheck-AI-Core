from pathlib import Path
from uuid import uuid4

from docx import Document
from fastapi.testclient import TestClient

from controlcheck.api import create_app
from controlcheck.storage import LocalFileStorage
from controlcheck.storage_s3 import S3FileStorage


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"


def _docx_text(path: Path) -> str:
    document = Document(path)
    blocks = [paragraph.text for paragraph in document.paragraphs]
    blocks.extend(cell.text for table in document.tables for row in table.rows for cell in row.cells)
    return "\n".join(blocks)


def test_health_live_probe():
    app = create_app()
    client = TestClient(app)
    response = client.get("/health/live")
    assert response.status_code == 200
    assert response.json()["status"] == "live"


def test_health_ready_probe_offline_mode():
    app = create_app()
    client = TestClient(app)
    response = client.get("/health/ready")
    assert response.status_code == 200
    data = response.json()
    assert data["status"] == "ready"
    assert "checks" in data
    assert data["checks"]["catalogue"] == "loaded"


def test_health_ready_does_not_expose_database_failure_details(tmp_path):
    def unavailable_session():
        raise RuntimeError("postgresql://operator:secret@internal/database")

    app = create_app(
        session_factory=unavailable_session,
        storage=LocalFileStorage(tmp_path),
    )
    client = TestClient(app)

    response = client.get("/health/ready")

    assert response.status_code == 503
    assert response.json()["status"] == "not_ready"
    assert response.json()["checks"]["database"] == "unreachable"
    assert "secret" not in response.text
    assert "postgresql" not in response.text


def test_health_ready_returns_503_when_storage_adapter_is_unavailable():
    class UnavailableStorage:
        def is_ready(self) -> bool:
            return False

    app = create_app(storage=UnavailableStorage())
    client = TestClient(app)

    response = client.get("/health/ready")

    assert response.status_code == 503
    assert response.json()["status"] == "not_ready"
    assert response.json()["checks"]["storage"] == "unavailable"


def test_security_headers_present():
    app = create_app()
    client = TestClient(app)
    response = client.get("/health/live")
    assert response.status_code == 200
    assert response.headers.get("X-Content-Type-Options") == "nosniff"
    assert response.headers.get("X-Frame-Options") == "DENY"
    assert response.headers.get("X-XSS-Protection") == "1; mode=block"
    assert response.headers.get("Referrer-Policy") == "strict-origin-when-cross-origin"
    assert "X-Request-ID" in response.headers


def test_prometheus_metrics_endpoint():
    app = create_app()
    client = TestClient(app)
    # Trigger a request to record metric
    client.get("/health/live")
    response = client.get("/metrics")
    assert response.status_code == 200
    assert "controlcheck_app_info" in response.text
    assert "controlcheck_http_requests_total" in response.text
    assert "controlcheck_process_uptime_seconds" in response.text


def test_production_settings_validation(monkeypatch):
    from controlcheck.settings import ProductionSettings
    import pytest

    # Test rejection of insecure JWT secret in production mode
    monkeypatch.setenv("CONTROLCHECK_ENV", "production")
    monkeypatch.setenv("CONTROLCHECK_JWT_SECRET", "dev-secret-key-change-in-production")
    with pytest.raises(ValueError, match="INSECURE CONFIGURATION"):
        ProductionSettings.from_env()

    # Test acceptance of secure 32+ char secret in production mode
    monkeypatch.setenv("CONTROLCHECK_JWT_SECRET", "a" * 64)
    monkeypatch.setenv(
        "CONTROLCHECK_DATABASE_URL",
        "postgresql+psycopg://controlcheck:controlcheck@database/controlcheck",
    )
    monkeypatch.setenv("CONTROLCHECK_CORS_ORIGINS", "https://controlcheck.example")
    monkeypatch.setenv("CONTROLCHECK_TRUSTED_HOSTS", "controlcheck.example")
    monkeypatch.setenv("CONTROLCHECK_STORAGE_BACKEND", "local")
    settings = ProductionSettings.from_env()
    assert settings.env == "production"
    assert len(settings.jwt_secret) == 64


def test_s3_storage_initialization():
    storage = S3FileStorage(bucket="my-test-bucket", region="ap-southeast-1")
    assert storage.bucket == "my-test-bucket"
    assert storage.region == "ap-southeast-1"


def test_prd_v08_records_production_readiness():
    text = _docx_text(DOCS / "ControlCheck_AI_PRD_v0.8.docx")
    assert "Product Requirements Document v0.8" in text
    assert "Phase 7 Production Readiness Alignment" in text
    assert "Change Log v0.8" in text
    assert "Multi-stage production Dockerfile" in text

