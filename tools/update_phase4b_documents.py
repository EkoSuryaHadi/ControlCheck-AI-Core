"""Generate the versioned Phase 4B specification documents."""
from __future__ import annotations

from datetime import datetime, timezone
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
STAMP = datetime(2026, 8, 24, tzinfo=timezone.utc)


def _doc(title: str, subtitle: str, sections: list[tuple[str, list[str]]], target: Path) -> None:
    document = Document()
    document.core_properties.title = title
    document.core_properties.modified = STAMP
    document.add_heading(title, 0)
    document.add_paragraph(subtitle)
    for heading, paragraphs in sections:
        document.add_heading(heading, level=1)
        for text in paragraphs:
            p = document.add_paragraph(text, style="List Bullet" if text.startswith("• ") else None)
            if text.startswith("• "):
                p.text = text[2:]
            p.paragraph_format.space_after = Pt(5)
    document.save(target)


def main() -> None:
    DOCS.mkdir(exist_ok=True)
    _doc("ERD & Database Specification v0.3", "ControlCheck AI | Phase 4B canonical ingestion | 24 August 2026", [
        ("Executable authority", ["Alembic migrations are the executable schema authority. The SQL file in docs/ is a readable reference and must not be applied directly."]),
        ("Canonical persistence", ["Immutable dataset snapshots preserve the uploaded source file, workbook SHA-256, source_project_name, data date, row counts, domain status, and mapping profile.", "Raw rows use BIGINT identity and retain source sheet, source row number, raw payload, and source key. Canonical WBS, budget, actual cost, commitment, schedule, and progress facts are snapshot-scoped."]),
        ("Governance contracts", ["The governed template is the only supported ingestion shape. Duplicate business IDs are retained as canonical facts while source_key persistence identity prevents collisions.", "Progress values above 100% and contradictory actual dates are stored unchanged; deterministic rules PRG-002 and DQ-003 report them.", "Partial-domain execution is explicit: rules run only when required domains are valid; blocked domains produce durable skipped-rule records.", "Authentication and complete RBAC remain deferred. X-Organization-ID is a controlled development tenant contract."]),
        ("Phase 4B tables", ["dataset_snapshots, dataset_domain_statuses, raw_rows, wbs_facts, budget_facts, actual_cost_facts, commitment_facts, schedule_facts, and progress_facts extend Phase 4A durable analysis tables."]),
    ], DOCS / "ControlCheck_AI_ERD_Database_Spec_v0.3.docx")
    _doc("Control Rule Catalogue v0.3", "ControlCheck AI | Deterministic 20-rule catalogue | 24 August 2026", [
        ("Runtime contract", ["All 20 MVP rules remain deterministic and catalogue-driven. Thresholds, severity, evidence requirements, and required domains are versioned with the catalogue SHA-256.", "A run records executed_rule_ids and skipped_rules. A skipped rule is not a finding and carries blocked_required_domain evidence for auditability."]),
        ("Data-quality semantics", ["DQ-002 treats duplicate business IDs as a finding while canonical persistence keeps each source row addressable with source_key. DQ-003 reports contradictory actual dates without rewriting source values."]),
        ("Progress and evidence", ["PRG-002 reports progress above 100% from the unchanged stored value. Every finding evidence item includes source sheet/rows, record IDs, fields, and BIGINT raw_row_ids when executed from a snapshot."]),
        ("AI boundary", ["LLM reasoning may summarize and prioritize deterministic findings, but cannot change rule calculations, thresholds, severity, or evidence."]),
    ], DOCS / "ControlCheck_AI_Control_Rule_Catalogue_v0.3.docx")
    _doc("Product Requirements Document v0.4", "ControlCheck AI | Canonical ingestion and database-native analysis | 24 August 2026", [
        ("Objective", ["Deliver a reproducible project-control audit workflow from governed workbook upload to immutable snapshot, canonical facts, deterministic analysis, traceable evidence, and durable API results."]),
        ("Phase 4B scope", ["PostgreSQL canonical ingestion for 149 governed rows across Project_Info (12), WBS (9), Budget (73), Actual_Cost (6), Commitments (13), Schedule (36), and Progress (not counted in source total when absent in the fixture).", "Snapshot upload is idempotent by project, filename, workbook SHA-256, mapping profile, and catalogue-compatible input; force_new creates a new snapshot for the same bytes.", "Golden parity acceptance: 59 findings, 100% precision and recall, and exact equality of deterministic finding payloads between Excel and database snapshot execution. Boundary acceptance: zero findings."]),
        ("Explicit semantics", ["Template-only ingestion; partial-domain execution; unchanged source anomalies; lossless source_project_name; BIGINT raw-row lineage; compatibility route retained for controlled migration." ]),
        ("Deferred", ["Authentication, complete RBAC, frontend, and LLM reasoning orchestration remain later phases. The current API is a backend validation surface, not a public production security boundary."]),
        ("Change log", ["v0.4 — Phase 4B canonical ingestion, immutable snapshots, domain gating, API snapshot resources, Excel-vs-DB parity, and updated database/rule governance."]),
    ], DOCS / "ControlCheck_AI_PRD_v0.4.docx")


if __name__ == "__main__":
    main()
