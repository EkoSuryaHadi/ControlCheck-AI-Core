"""Generate the immutable ControlCheck internal-production PRD v0.5."""
from __future__ import annotations

from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
TARGET = ROOT / "docs" / "ControlCheck_AI_PRD_v0.5.docx"
STAMP = datetime(2026, 8, 24, tzinfo=timezone.utc)
BLUE = "2E74B5"
DARK_BLUE = "0B2545"
MUTED = "5B6573"
LIGHT_FILL = "F2F4F7"
CALLOUT_FILL = "E8EEF5"
WHITE = "FFFFFF"
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def set_run_font(
    run,
    *,
    name: str = "Calibri",
    size: float | None = None,
    color: str | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_repeat_table_header(row) -> None:
    properties = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    properties.append(header)


def set_cell_shading(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def shade_paragraph(paragraph, fill: str) -> None:
    properties = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)
    borders = OxmlElement("w:pBdr")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "8")
        border.set(qn("w:color"), "B9C8D8")
        borders.append(border)
    properties.append(borders)


def configure_table_geometry(table, widths_dxa: tuple[int, ...]) -> None:
    if sum(widths_dxa) != TABLE_WIDTH_DXA:
        raise ValueError("Table column widths must total 9360 DXA")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    properties = table._tbl.tblPr

    width = properties.find(qn("w:tblW"))
    if width is None:
        width = OxmlElement("w:tblW")
        properties.append(width)
    width.set(qn("w:type"), "dxa")
    width.set(qn("w:w"), str(TABLE_WIDTH_DXA))

    indent = properties.find(qn("w:tblInd"))
    if indent is None:
        indent = OxmlElement("w:tblInd")
        properties.append(indent)
    indent.set(qn("w:type"), "dxa")
    indent.set(qn("w:w"), str(TABLE_INDENT_DXA))

    layout = properties.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        properties.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for value in widths_dxa:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(value))
        grid.append(column)

    for row in table.rows:
        cant_split = OxmlElement("w:cantSplit")
        row._tr.get_or_add_trPr().append(cant_split)
        for index, cell in enumerate(row.cells):
            value = widths_dxa[index]
            cell.width = Inches(value / 1440)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell_properties = cell._tc.get_or_add_tcPr()
            cell_width = cell_properties.find(qn("w:tcW"))
            if cell_width is None:
                cell_width = OxmlElement("w:tcW")
                cell_properties.append(cell_width)
            cell_width.set(qn("w:type"), "dxa")
            cell_width.set(qn("w:w"), str(value))
            margins = cell_properties.find(qn("w:tcMar"))
            if margins is None:
                margins = OxmlElement("w:tcMar")
                cell_properties.append(margins)
            for side, margin_value in (
                ("top", 80),
                ("bottom", 80),
                ("start", 120),
                ("end", 120),
            ):
                margin = margins.find(qn(f"w:{side}"))
                if margin is None:
                    margin = OxmlElement(f"w:{side}")
                    margins.append(margin)
                margin.set(qn("w:w"), str(margin_value))
                margin.set(qn("w:type"), "dxa")


def add_table(document, headers, rows, widths_dxa):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, text in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = text
        set_cell_shading(cell, LIGHT_FILL)
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        for run in paragraph.runs:
            set_run_font(run, size=9.5, color=DARK_BLUE, bold=True)
    set_repeat_table_header(table.rows[0])

    for values in rows:
        cells = table.add_row().cells
        for index, text in enumerate(values):
            cells[index].text = text
            paragraph = cells[index].paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.08
            for run in paragraph.runs:
                set_run_font(run, size=9.25, color="20242A")
    configure_table_geometry(table, tuple(widths_dxa))
    after = document.add_paragraph()
    after.paragraph_format.space_after = Pt(2)
    return table


def add_bullet(document, text: str) -> None:
    paragraph = document.add_paragraph(text, style="List Bullet")
    paragraph.paragraph_format.keep_together = True


def add_number(document, text: str) -> None:
    paragraph = document.add_paragraph(text, style="List Number")
    paragraph.paragraph_format.keep_together = True


def add_heading(document, text: str, level: int = 1) -> None:
    paragraph = document.add_heading(text, level=level)
    paragraph.paragraph_format.keep_with_next = True


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, instruction, separate, text, end):
        run._r.append(element)
    set_run_font(run, size=9, color=MUTED)


def configure_styles(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167


def configure_page(document: Document) -> None:
    # Define both variants explicitly. Word normally mirrors the primary
    # header/footer when odd/even mode is disabled, but LibreOffice can retain
    # blank or partial even-page parts during headless conversion.
    document.settings.odd_and_even_pages_header_footer = True
    section = document.sections[0]
    section.different_first_page_header_footer = False
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    for header_part in (section.header, section.even_page_header):
        header = header_part.paragraphs[0]
        header.text = "CONTROLCHECK AI  |  PRODUCT REQUIREMENTS"
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in header.runs:
            set_run_font(run, size=8.5, color=MUTED, bold=True)

    for footer_part in (section.footer, section.even_page_footer):
        footer = footer_part.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        prefix = footer.add_run("PRD v0.5  |  24 August 2026  |  Page ")
        set_run_font(prefix, size=9, color=MUTED)
        add_page_field(footer)


def add_masthead(document: Document) -> None:
    kicker = document.add_paragraph()
    kicker.paragraph_format.space_after = Pt(4)
    set_run_font(
        kicker.add_run("INTERNAL PRODUCTION MVP"),
        size=10,
        color=BLUE,
        bold=True,
    )

    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    set_run_font(
        title.add_run("Product Requirements Document v0.5"),
        size=24,
        color=DARK_BLUE,
        bold=True,
    )

    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(16)
    set_run_font(
        subtitle.add_run("ControlCheck Core - secure single-organization pilot"),
        size=14,
        color=MUTED,
    )

    for label, value in (
        ("Status", "Approved implementation baseline"),
        ("Release", "Internal API-only production MVP"),
        ("Date", "24 August 2026"),
        ("Executable authority", "Application code, runtime JSON catalogues, and Alembic migrations"),
    ):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(2)
        set_run_font(paragraph.add_run(f"{label}: "), size=10.5, bold=True)
        set_run_font(paragraph.add_run(value), size=10.5, color="20242A")

    callout = document.add_paragraph()
    callout.paragraph_format.left_indent = Inches(0.08)
    callout.paragraph_format.right_indent = Inches(0.08)
    callout.paragraph_format.space_before = Pt(12)
    callout.paragraph_format.space_after = Pt(12)
    shade_paragraph(callout, CALLOUT_FILL)
    run = callout.add_run(
        "Decision: ship one secure internal API replica with bearer-key authentication, "
        "a server-fixed tenant, managed PostgreSQL, and persistent workbook storage."
    )
    set_run_font(run, size=10.5, color=DARK_BLUE, bold=True)


def build_document() -> Document:
    document = Document()
    document.core_properties.title = "ControlCheck AI Product Requirements Document v0.5"
    document.core_properties.subject = "Internal production MVP"
    document.core_properties.author = "ControlCheck AI"
    document.core_properties.keywords = "ControlCheck, production, API, project controls"
    document.core_properties.created = STAMP
    document.core_properties.modified = STAMP
    configure_styles(document)
    configure_page(document)
    add_masthead(document)

    add_heading(document, "1. Executive summary")
    document.add_paragraph(
        "ControlCheck Core converts governed EPC project-control workbooks into immutable "
        "snapshots, canonical facts, deterministic findings, and traceable evidence. PRD "
        "v0.5 promotes the validated Phase 4B backend into a deliberately narrow internal "
        "production pilot without changing the 20-rule calculation boundary."
    )
    document.add_paragraph(
        "The pilot is secure enough for one controlled organization: all versioned API "
        "routes require a bearer API key, tenant identity is fixed server-side, health "
        "responses disclose minimal state, and release gates preserve Golden 59 and "
        "Boundary zero behavior."
    )

    add_heading(document, "2. Product objective and success")
    document.add_paragraph(
        "Enable an internal project-controls team to deploy and operate the deterministic "
        "ControlCheck workflow on managed infrastructure, while preserving reproducibility, "
        "tenant isolation, evidence lineage, and recoverability."
    )
    add_bullet(document, "A reviewed image starts predictably from environment-only configuration.")
    add_bullet(document, "Unauthorized callers cannot access any /v1/* operation.")
    add_bullet(document, "The client cannot select or override the production organization.")
    add_bullet(document, "Database, upload storage, and catalogue readiness are observable without leaking internals.")
    add_bullet(document, "Controlled fixtures remain Golden 59 and Boundary zero under strict evaluation.")

    add_heading(document, "3. Release scope")
    add_table(
        document,
        ("Area", "In v0.5", "Explicitly deferred"),
        (
            ("Identity", "One bearer API key; one fixed organization UUID", "JWT/OIDC, users, roles, complete RBAC"),
            ("Runtime", "One non-root container replica", "Horizontal scaling and rolling multi-replica writes"),
            ("Data", "Managed PostgreSQL plus mounted persistent workbook volume", "S3-compatible object storage and cross-region replication"),
            ("Product", "API-only deterministic analysis and durable query workflow", "Frontend, customer self-service, LLM orchestration"),
            ("Tenancy", "One controlled internal organization", "Public multi-tenant SaaS onboarding and billing"),
        ),
        (1800, 3780, 3780),
    )

    add_heading(document, "4. Users and operating model")
    add_heading(document, "4.1 Internal analyst", 2)
    document.add_paragraph(
        "Submits governed workbooks, starts snapshot analysis, reviews findings and evidence, "
        "and updates finding status through an approved internal client."
    )
    add_heading(document, "4.2 Production operator", 2)
    document.add_paragraph(
        "Controls deployment settings, secrets, migrations, backups, restores, API key "
        "rotation, image promotion, rollback, and request-ID-led incident diagnostics."
    )
    add_heading(document, "4.3 Engineering reviewer", 2)
    document.add_paragraph(
        "Approves pull requests only after database, tests, deterministic fixtures, document "
        "contracts, and container build gates succeed."
    )

    add_heading(document, "5. Functional requirements")
    add_heading(document, "5.1 Runtime configuration", 2)
    add_bullet(document, "Production mode is selected only by CONTROLCHECK_ENV=production.")
    add_bullet(document, "Startup fails when database URL, absolute upload root, readable catalogue, organization UUID, API key, or exact trusted hosts are invalid or missing.")
    add_bullet(document, "The API key contains at least 32 characters; maximum upload size is a positive byte count.")
    add_bullet(document, "CORS defaults to absent and accepts only explicitly configured exact HTTPS origins.")
    add_bullet(document, "Interactive docs and OpenAPI are disabled by default in production.")

    add_heading(document, "5.2 Authentication and tenant isolation", 2)
    add_bullet(document, "Every /v1/* route requires Authorization: Bearer <key> in production.")
    add_bullet(document, "Missing, malformed, or incorrect credentials return a generic 401 error and WWW-Authenticate: Bearer.")
    add_bullet(document, "Credential comparison is constant-time and credentials never appear in committed files or response bodies.")
    add_bullet(document, "Production tenant identity comes only from CONTROLCHECK_ORGANIZATION_ID; X-Organization-ID is ignored.")
    add_bullet(document, "Development retains the tenant header solely as a controlled non-production contract.")

    add_heading(document, "5.3 Deterministic workflow", 2)
    add_bullet(document, "Preserve the Phase 4B governed-template, immutable-snapshot, raw-lineage, canonical-fact, domain-gating, and durable-evidence contracts.")
    add_bullet(document, "Keep calculations, thresholds, severities, finding metrics, and evidence deterministic and independent from any LLM.")
    add_bullet(document, "Retain the compatibility workbook audit endpoint during the internal pilot.")
    add_bullet(document, "Enforce organization scope on project, snapshot, run, finding, evidence, and status operations.")

    add_heading(document, "5.4 Health and safe HTTP behavior", 2)
    add_bullet(document, "GET /health/live returns only status=live when the process can answer HTTP.")
    add_bullet(document, "GET /health/ready checks database, persistent storage, and catalogue, returning only ready or not_ready.")
    add_bullet(document, "The legacy /health response remains during the pilot for compatibility.")
    add_bullet(document, "Production accepts only configured hostnames and validates or replaces caller-supplied request IDs.")
    add_bullet(document, "Unhandled errors return a generic envelope with a safe request ID and no traceback or dependency detail.")

    add_heading(document, "6. API surface")
    add_table(
        document,
        ("Surface", "Purpose", "Production access"),
        (
            ("/health/live", "Process liveness", "Public; minimal"),
            ("/health/ready", "Database, storage, and catalogue readiness", "Public; minimal"),
            ("/health", "Pilot compatibility health", "Public"),
            ("/v1/audits", "Stateless governed workbook audit", "Bearer key"),
            ("/v1/organizations/.../projects", "Organization-scoped project lifecycle", "Bearer key + fixed tenant"),
            ("/v1/projects/.../dataset-snapshots", "Snapshot ingestion and retrieval", "Bearer key + fixed tenant"),
            ("/v1/.../analysis-runs", "Deterministic analysis and history", "Bearer key + fixed tenant"),
            ("/v1/findings/...", "Finding, evidence, filters, and status", "Bearer key + fixed tenant"),
        ),
        (2500, 4160, 2700),
    )

    add_heading(document, "7. Deployment and release requirements")
    add_number(document, "Build from Python 3.11.9 slim-bookworm and run the application as UID/GID 10001.")
    add_number(document, "Apply the single linear Alembic head before Uvicorn starts; production database URL comes from environment configuration.")
    add_number(document, "Mount one persistent upload volume at the configured absolute path and operate exactly one replica.")
    add_number(document, "Use managed PostgreSQL 16 with TLS, automated backup, restore testing, and credentials held outside Git.")
    add_number(document, "Require CI compilation, migrations, Alembic drift check, full pytest, Golden strict, Boundary strict, and Docker build.")
    add_number(document, "Promote an immutable reviewed image digest from staging to production.")

    add_heading(document, "8. Security and non-functional requirements")
    add_table(
        document,
        ("Requirement", "Acceptance"),
        (
            ("Confidentiality", "No real API keys, database credentials, organization IDs, or private keys in Git, images, logs, or responses."),
            ("Isolation", "All durable repository access remains organization-scoped; client headers cannot change the production tenant."),
            ("Availability", "Live and ready probes distinguish process health from dependency readiness; failed readiness returns 503."),
            ("Recoverability", "Database and workbook-volume backups are restored together and verified through persisted evidence retrieval."),
            ("Traceability", "Every response carries a safe request ID; deterministic findings retain calculation and source evidence."),
            ("Performance", "Upload size is bounded at 25 MiB by default; the pilot uses one replica and controlled internal load."),
            ("Maintainability", "Configuration is validated once, migrations remain linear, and each behavior change includes a regression test."),
        ),
        (2200, 7160),
    )

    add_heading(document, "9. Acceptance criteria")
    add_table(
        document,
        ("Gate", "Required result", "Evidence"),
        (
            ("Configuration", "Invalid production settings fail startup", "Production settings tests"),
            ("Authentication", "All /v1/* routes reject missing/wrong key with 401", "Route enumeration and HTTP tests"),
            ("Tenant", "Client tenant header cannot override configured UUID", "Fixed-tenant contract test"),
            ("HTTP safety", "Docs off, host restricted, CORS opt-in, generic 500", "Production security tests"),
            ("Health", "Live 200; ready 200 or minimal 503", "Dependency readiness tests"),
            ("Container", "Build succeeds; runtime UID 10001; no reload", "Image build and inspection"),
            ("Database", "Exactly one head; upgrade and drift check succeed", "Alembic CI gate"),
            ("Engine", "Golden 59 with zero FP/FN; Boundary zero", "Strict fixture evaluations"),
            ("Operations", "Rotation, restore verification, rollback, and incidents documented", "Production runbook"),
        ),
        (1900, 4100, 3360),
    )

    add_heading(document, "10. Risks and controls")
    add_table(
        document,
        ("Risk", "Control", "Residual decision"),
        (
            ("Shared bearer key has no user attribution", "Restricted secret distribution, request IDs, key rotation, internal-only access", "Accept for pilot; replace with JWT/RBAC later"),
            ("Local volume prevents safe horizontal writes", "Exactly one replica, persistent mount, paired backup verification", "Accept for pilot; move to object storage before scaling"),
            ("Migration failure blocks startup", "CI migration gates, managed backup, reviewed rollback/forward-fix procedure", "Fail closed"),
            ("Controlled fixtures overstate real-world confidence", "Label results as synthetic agreement; require customer validation before accuracy claims", "No customer accuracy claim"),
            ("Public readiness can aid reconnaissance", "Minimal status only; no dependency names or exceptions", "Accept behind HTTPS ingress"),
        ),
        (2700, 4260, 2400),
    )

    add_heading(document, "11. Deferred roadmap")
    add_bullet(document, "JWT/OIDC authentication, named users, role model, and complete RBAC.")
    add_bullet(document, "S3-compatible object storage, multi-replica execution, and scalable background jobs.")
    add_bullet(document, "Production frontend, customer onboarding, subscription/billing, and public multi-tenancy.")
    add_bullet(document, "LLM reasoning that summarizes deterministic finding and evidence payloads without altering rule outcomes.")
    add_bullet(document, "Customer-data validation, operational SLOs, and capacity targets based on observed pilot load.")

    add_heading(document, "12. Change log")
    document.add_paragraph(
        "v0.5 - Defines the internal production MVP security boundary, strict runtime "
        "configuration, bearer authentication, fixed tenant, safe HTTP behavior, liveness "
        "and readiness, managed PostgreSQL, persistent storage, non-root container, CI release "
        "gates, operating runbook, and explicit JWT/RBAC/object-storage/frontend deferrals."
    )
    document.add_paragraph(
        "Historical PRDs remain immutable. PRD v0.4 continues to describe canonical ingestion; "
        "v0.5 adds the production operating boundary without replacing runtime JSON catalogues "
        "or executable Alembic migrations."
    )
    return document


def main() -> None:
    TARGET.parent.mkdir(parents=True, exist_ok=True)
    document = build_document()
    document.save(TARGET)
    print(TARGET)


if __name__ == "__main__":
    main()
