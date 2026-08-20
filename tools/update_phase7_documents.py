"""Create versioned Phase 7 specification documents without mutating history."""

from __future__ import annotations

from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.text import WD_BREAK
from docx.shared import Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
FIXED_MODIFIED = datetime(2026, 8, 20, tzinfo=timezone.utc)


def _replace(document: DocumentObject, old: str, new: str) -> None:
    for paragraph in document.paragraphs:
        if paragraph.text.strip() == old:
            paragraph.text = new
            return
    raise ValueError(f"Source heading not found: {old}")


def _replace_table_row(document: DocumentObject, key: str, replacements: dict[int, str]) -> None:
    for table in document.tables:
        for row in table.rows:
            if row.cells and row.cells[0].text.strip() == key:
                for index, value in replacements.items():
                    row.cells[index].text = value
                return
    raise ValueError(f"Source table row not found: {key}")


def _heading(document: DocumentObject, text: str, level: int = 1) -> None:
    paragraph = document.add_heading(text, level=level)
    paragraph.paragraph_format.keep_with_next = True


def _paragraph(document: DocumentObject, text: str, *, bullet: bool = False) -> None:
    style = "List Bullet" if bullet else None
    paragraph = document.add_paragraph(text, style=style)
    paragraph.paragraph_format.space_after = Pt(5)


def _table(document: DocumentObject, headers: tuple[str, ...], rows: list[tuple[str, ...]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True
    for cell, value in zip(table.rows[0].cells, headers, strict=True):
        cell.text = value
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
        cell._tc.get_or_add_tcPr().append(_shading("123047"))
    for row_values in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row_values, strict=True):
            cell.text = value
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)
    document.add_paragraph()


def _shading(fill: str):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    return shading


def _prepare(source_name: str, target_name: str, title: str, old_title: str) -> DocumentObject:
    document = Document(DOCS / source_name)
    _replace(document, old_title, title)
    document.core_properties.title = title
    document.core_properties.modified = FIXED_MODIFIED
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    return document


def build_prd() -> Path:
    target = DOCS / "ControlCheck_AI_PRD_v0.8.docx"
    document = _prepare(
        "ControlCheck_AI_PRD_v0.7.docx",
        target.name,
        "Product Requirements Document v0.8",
        "Product Requirements Document v0.7",
    )
    _replace(document, "Version 0.7 | AI Intelligence Layer & Safety Guardrails | 20 August 2026", "Version 0.8 | Production Readiness & Deployment Engineering | 20 August 2026")
    _replace_table_row(document, "Version", {1: "0.8"})
    _heading(document, "Phase 7 Production Readiness Alignment", 1)
    _paragraph(document, "Phase 7 establishes production hardening and deployment engineering standards for ControlCheck AI, ensuring container security, database resilience, cloud storage options, and automated CI/CD validation.")
    _heading(document, "Scope delivered in Phase 7", 2)
    for item in (
        "Multi-stage production Dockerfile with non-root security context (UID 10001).",
        "Production Docker Compose setup (docker-compose.prod.yml) with PostgreSQL health checks.",
        "Container liveness (/health/live) and database readiness (/health/ready) probes.",
        "CORS middleware with configurable domain whitelist (CONTROLCHECK_CORS_ORIGINS).",
        "S3 / MinIO cloud object storage backend implementation (S3FileStorage).",
        "Automated GitHub Actions CI pipeline (.github/workflows/ci.yml).",
        "Comprehensive Production Runbook (docs/PRODUCTION_RUNBOOK.md).",
    ):
        _paragraph(document, item, bullet=True)
    _heading(document, "Change Log v0.8", 2)
    _table(
        document,
        ("Version", "Date", "Change"),
        [
            ("0.1", "17 Aug 2026", "Initial MVP blueprint."),
            ("0.2", "17 Aug 2026", "Rule and synthetic validation alignment."),
            ("0.3", "17 Aug 2026", "Phase 4A PostgreSQL persistence, durable API, storage, lifecycle, and Phase 4B sequencing."),
            ("0.4", "20 Aug 2026", "Phase 4B raw-row lineage, canonical fact normalization, and import batch tracking."),
            ("0.5", "20 Aug 2026", "Phase 5A structured logging, Phase 5B pagination & idempotency, Phase 5C health scoring engine."),
            ("0.6", "20 Aug 2026", "Phase 4C JWT authentication, bcrypt password hashing, and RBAC membership controls."),
            ("0.7", "20 Aug 2026", "Phase 6 AI Intelligence Layer, safety guardrails, controlled tools, and audit assistant endpoints."),
            ("0.8", "20 Aug 2026", "Phase 7 Production Readiness, containerization, CORS hardening, probes, CI/CD, and runbook."),
        ],
    )
    document.save(target)
    return target


def main() -> None:
    output = build_prd()
    print("Generated:", output)


if __name__ == "__main__":
    main()
