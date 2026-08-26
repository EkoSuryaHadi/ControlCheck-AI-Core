"""Create the versioned public-beta PRD without mutating prior PRDs."""

from __future__ import annotations

from datetime import datetime, timezone
import os
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile, ZipInfo

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.text import WD_BREAK
from docx.shared import Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
FIXED_MODIFIED = datetime(2026, 8, 26, tzinfo=timezone.utc)
FIXED_ZIP_TIMESTAMP = (2026, 8, 26, 0, 0, 0)


def _replace(document: DocumentObject, old: str, new: str) -> None:
    for paragraph in document.paragraphs:
        if paragraph.text.strip() == old:
            paragraph.text = new
            return
    raise ValueError(f"Source text not found: {old}")


def _replace_table_row(document: DocumentObject, key: str, replacements: dict[int, str]) -> None:
    for table in document.tables:
        for row in table.rows:
            if row.cells and row.cells[0].text.strip() == key:
                for index, value in replacements.items():
                    row.cells[index].text = value
                return
    raise ValueError(f"Source table row not found: {key}")


def _truncate_from_paragraph(document: DocumentObject, marker: str) -> None:
    """Remove a superseded trailing amendment while preserving section settings."""
    body = document._element.body
    marker_element = next(
        (paragraph._element for paragraph in document.paragraphs if paragraph.text.strip() == marker),
        None,
    )
    if marker_element is None:
        raise ValueError(f"Source text not found: {marker}")
    removing = False
    for element in list(body):
        if element is marker_element:
            removing = True
        if removing and element.tag.rsplit("}", 1)[-1] != "sectPr":
            body.remove(element)


def _heading(document: DocumentObject, text: str, level: int = 1) -> None:
    paragraph = document.add_heading(text, level=level)
    paragraph.paragraph_format.keep_with_next = True


def _paragraph(document: DocumentObject, text: str, *, bullet: bool = False) -> None:
    paragraph = document.add_paragraph(text, style="List Bullet" if bullet else None)
    paragraph.paragraph_format.space_after = Pt(2)


def _shading(fill: str):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    return shading


def _table(document: DocumentObject, headers: tuple[str, ...], rows: list[tuple[str, ...]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True
    for cell, value in zip(table.rows[0].cells, headers, strict=True):
        cell.text = value
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
        cell._tc.get_or_add_tcPr().append(_shading("123047"))
    for row_values in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row_values, strict=True):
            cell.text = value
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)
    document.add_paragraph()


def _normalize_docx_archive(path: Path) -> None:
    """Rewrite the DOCX archive with a fixed member order and metadata."""
    normalized_path = path.with_name(f"{path.stem}.normalized{path.suffix}")
    with ZipFile(path, "r") as source, ZipFile(
        normalized_path,
        "w",
        compression=ZIP_DEFLATED,
        compresslevel=9,
    ) as target:
        target.comment = b""
        for source_info in sorted(source.infolist(), key=lambda item: item.filename):
            target_info = ZipInfo(source_info.filename, date_time=FIXED_ZIP_TIMESTAMP)
            target_info.compress_type = ZIP_DEFLATED
            target_info.create_system = 0
            target_info.create_version = 20
            target_info.extract_version = 20
            target_info.flag_bits = 0
            target_info.internal_attr = 0
            target_info.external_attr = 0
            target.writestr(
                target_info,
                source.read(source_info.filename),
                compress_type=ZIP_DEFLATED,
                compresslevel=9,
            )
    os.replace(normalized_path, path)


def build_prd() -> Path:
    target = DOCS / "ControlCheck_AI_PRD_v1.2.docx"
    document = Document(DOCS / "ControlCheck_AI_PRD_v1.1.docx")
    _replace(document, "Product Requirements Document v1.1", "Product Requirements Document v1.2")
    _replace(
        document,
        "Version 1.1 | Public Beta Cloud Deployment & Usage Validation | 26 August 2026",
        "Version 1.2 | Vercel Hybrid Public Beta Deployment | 26 August 2026",
    )
    _replace_table_row(document, "Version", {1: "1.2"})
    _truncate_from_paragraph(
        document,
        "Phase 10 Public Beta Cloud Deployment & Usage Validation Alignment",
    )
    document.core_properties.title = "Product Requirements Document v1.2"
    document.core_properties.modified = FIXED_MODIFIED
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    _heading(document, "Phase 10 Vercel Hybrid Public Beta Deployment Alignment")
    _paragraph(
        document,
        "Vercel React + FastAPI is the active no-card public-beta deployment. The topology is browser → Vercel React frontend → Vercel FastAPI Function → Supabase PostgreSQL + private Cloudflare R2.",
    )
    _heading(document, "Hosted data flow", 2)
    _paragraph(
        document,
        "register/login → create project → upload workbook → persist workbook in R2 → canonical ingestion and deterministic analysis in the Vercel FastAPI Function → persist run/findings/evidence in Supabase → display results in Vercel.",
    )
    _heading(document, "Runtime boundaries and release controls", 2)
    for item in (
        "Workbook uploads are limited to 4 MiB for the initial public beta because the Vercel request payload limit is 4.5 MB.",
        "The FastAPI Function has a 240-second configured duration target and its packaged runtime must remain below the 500 MB uncompressed bundle limit.",
        "Database migrations are an explicit release step before deployment; they never run implicitly during serverless startup.",
        "Production must fail closed when database, durable R2 storage, catalogue, JWT, CORS, or trusted-host configuration is missing or invalid.",
        "The no-card topology avoids a paid always-on backend while preserving durable PostgreSQL records and private workbook storage.",
        "A future direct-to-R2 upload path can remove the serverless request-body constraint without changing deterministic analysis semantics.",
    ):
        _paragraph(document, item, bullet=True)
    _heading(document, "Public-beta success criteria", 2)
    for item in (
        "hosted register-to-findings flow passes.",
        "uploaded files and results persist across Vercel Function invocations and deployments.",
        "no secrets appear in source/logs/docs.",
        "registrations, active users, projects, workbook uploads, and completed analysis runs are measurable from persisted records.",
    ):
        _paragraph(document, item, bullet=True)
    _heading(document, "Free-tier constraints and upgrade triggers", 2)
    for item in (
        "Vercel Function cold starts and execution limits may affect large analyses.",
        "Supabase Free may pause after low activity and has limited capacity/no managed downloadable backups.",
        "R2 Standard free allowance is used and the bucket remains private.",
        "Upgrade when user experience is materially affected by serverless limits, Supabase database approaches 400 MB, R2 approaches 8 GB, or usage becomes routine enough to justify an always-on backend.",
    ):
        _paragraph(document, item, bullet=True)
    _heading(document, "Deferred non-goals", 2)
    _paragraph(
        document,
        "full authentication/RBAC hardening, payment/subscription, enterprise SSO, and production-scale HA/DR remain deferred.",
    )
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    _heading(document, "Change Log v1.2", 2)
    _table(
        document,
        ("Version", "Date", "Change"),
        [
            (
                "1.2",
                "26 Aug 2026",
                "Vercel hybrid public-beta architecture, 4 MiB upload boundary, and explicit release controls.",
            )
        ],
    )
    document.save(target)
    _normalize_docx_archive(target)
    return target


def main() -> None:
    print("Generated:", build_prd())


if __name__ == "__main__":
    main()
