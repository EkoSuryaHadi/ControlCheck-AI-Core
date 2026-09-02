"""Create PRD v1.3 for AI Insight v1 without changing prior PRDs."""

from __future__ import annotations

from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.text import WD_BREAK
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"


def replace_exact(document: DocumentObject, old: str, new: str) -> None:
    for paragraph in document.paragraphs:
        if paragraph.text.strip() == old:
            paragraph.text = new
            return
    raise ValueError(f"Source text not found: {old}")


def heading(document: DocumentObject, text: str, level: int = 1) -> None:
    paragraph = document.add_heading(text, level=level)
    paragraph.paragraph_format.keep_with_next = True


def paragraph(document: DocumentObject, text: str, bullet: bool = False) -> None:
    item = document.add_paragraph(text, style="List Bullet" if bullet else None)
    item.paragraph_format.space_after = Pt(3)


def build_prd() -> Path:
    target = DOCS / "ControlCheck_AI_PRD_v1.3.docx"
    document = Document(DOCS / "ControlCheck_AI_PRD_v1.2.docx")
    replace_exact(document, "Product Requirements Document v1.2", "Product Requirements Document v1.3")
    replace_exact(
        document,
        "Version 1.2 | Vercel Hybrid Public Beta Deployment | 26 August 2026",
        "Version 1.3 | Evidence-Grounded AI Insight Public Beta | 2 September 2026",
    )
    document.core_properties.title = "Product Requirements Document v1.3"
    document.core_properties.modified = datetime(2026, 9, 2, tzinfo=timezone.utc)
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    heading(document, "AI Insight v1 Public Beta Alignment")
    paragraph(document, "AI Insight v1 generates a concise Indonesian project-control explanation after deterministic analysis completes. The deterministic engine remains the source of truth; AI only explains persisted findings and approved summaries.")
    heading(document, "Evidence and privacy contract", 2)
    for item in (
        "The insight input is bounded to severity/category aggregates, deduplicated top findings, recommendations, and finding identifiers from the selected analysis run.",
        "Workbook raw content is not sent to the AI provider. Source files, sheet rows, cell values, file paths, vendor rows, and telemetry payloads are excluded.",
        "The model must cite only persisted finding identifiers from the current tenant and analysis run, and must state missing cost or progress coverage as a limitation.",
        "Prompt text is treated as untrusted data and cannot override the server-side system instructions.",
    ):
        paragraph(document, item, bullet=True)
    heading(document, "Best-effort delivery and controls", 2)
    for item in (
        "A pending insight is persisted for each completed run before generation is attempted.",
        "OpenAI failures never change a completed deterministic analysis into a failed analysis.",
        "Only one ready insight is stored per analysis run; failed or pending insight generation can be retried through an authenticated endpoint.",
        "Vercel best-effort execution is the current beta delivery mechanism. Durable worker-based generation for large asynchronous imports remains deferred.",
    ):
        paragraph(document, item, bullet=True)
    heading(document, "Change Log v1.3", 2)
    paragraph(document, "2 September 2026 — AI Insight v1 added: evidence-grounded OpenAI summaries, tenant-scoped persistence, safe retry behavior, and dashboard states. Interactive AI chat remains out of scope.")
    document.save(target)
    return target


if __name__ == "__main__":
    print(build_prd())