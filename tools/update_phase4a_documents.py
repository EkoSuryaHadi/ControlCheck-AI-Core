"""Create versioned Phase 4A specification documents without mutating history."""

from __future__ import annotations

import json
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_BREAK
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
CATALOGUE_PATH = ROOT / "data" / "controlcheck_rule_catalogue_v0.2.json"
FIXED_MODIFIED = datetime(2026, 8, 17, tzinfo=timezone.utc)


def _replace(document: DocumentObject, old: str, new: str) -> None:
    for paragraph in document.paragraphs:
        if paragraph.text.strip() == old:
            paragraph.text = new
            return
    raise ValueError(f"Source heading not found: {old}")


def _replace_table_row(document: DocumentObject, key: str, replacements: dict[int, str]) -> None:
    for table in document.tables:
        for row in table.rows:
            if row.cells and row.cells[0].text.strip() == key:
                for index, value in replacements.items():
                    row.cells[index].text = value
                return
    raise ValueError(f"Source table row not found: {key}")


def _heading(document: DocumentObject, text: str, level: int = 1) -> None:
    paragraph = document.add_heading(text, level=level)
    paragraph.paragraph_format.keep_with_next = True


def _paragraph(document: DocumentObject, text: str, *, bullet: bool = False) -> None:
    style = "List Bullet" if bullet else None
    paragraph = document.add_paragraph(text, style=style)
    paragraph.paragraph_format.space_after = Pt(5)


def _table(document: DocumentObject, headers: tuple[str, ...], rows: list[tuple[str, ...]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True
    for cell, value in zip(table.rows[0].cells, headers, strict=True):
        cell.text = value
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
        cell._tc.get_or_add_tcPr().append(_shading("123047"))
    for row_values in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row_values, strict=True):
            cell.text = value
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)
    document.add_paragraph()


def _shading(fill: str):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    return shading


def _prepare(source_name: str, target_name: str, title: str, old_title: str) -> DocumentObject:
    document = Document(DOCS / source_name)
    _replace(document, old_title, title)
    document.core_properties.title = title
    document.core_properties.modified = FIXED_MODIFIED
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    return document


def build_erd() -> Path:
    target = DOCS / "ControlCheck_AI_ERD_Database_Spec_v0.2.docx"
    document = _prepare(
        "ControlCheck_AI_ERD_Database_Spec_v0.1.docx",
        target.name,
        "ERD & Database Specification v0.2",
        "ERD & Database Specification v0.1",
    )
    _heading(document, "Phase 4A Persistence Alignment", 1)
    _paragraph(document, "Version 0.2 aligns the conceptual ERD with the executable PostgreSQL persistence layer delivered in Phase 4A. Alembic migration 20260817_0001 is the executable schema authority; the SQL reference and this document are readable specifications.")
    _heading(document, "Implemented Phase 4A tables", 2)
    rows = [
        ("organizations", "Tenant boundary and lifecycle status", "Parent of all tenant-scoped records"),
        ("projects", "Project identity and planning envelope", "Unique organization + project code"),
        ("source_files", "Immutable workbook metadata and SHA-256", "Project-scoped storage reference"),
        ("dataset_snapshots", "Validated analysis input at a data date", "Connects project, source file, and source project ID"),
        ("rule_catalogue_versions", "Versioned JSONB rule definitions", "Version + SHA-256 unique identity"),
        ("analysis_runs", "Durable engine execution lifecycle", "queued/running/succeeded/failed"),
        ("findings", "Materialized deterministic findings", "Unique engine finding per analysis run"),
        ("finding_evidence", "Ordered, source-traceable evidence", "Rows, record IDs, fields, aggregation in JSONB"),
        ("approved_exceptions", "Governed suppressions and validity window", "Execution support is scheduled after Phase 4A"),
        ("audit_logs", "Append-oriented lifecycle event history", "Records status changes and tenant context"),
    ]
    _table(document, ("Table", "Purpose", "Key contract"), rows)
    _heading(document, "Contract decisions", 2)
    for item in (
        "All tenant-sensitive queries include organization_id in the database predicate.",
        "Project workbook identity must match the registered project code before a durable run is created.",
        "findings.entity_id is text (up to 300 characters), because deterministic engine entities include composite identifiers—not only UUIDs.",
        "A failed engine execution retains an analysis_runs record with a safe error code and message, but creates no findings.",
        "finding_evidence preserves ordered JSONB source references for auditability while raw and canonical fact tables remain Phase 4B scope.",
        "Authentication and complete RBAC are deferred; Phase 4A uses X-Organization-ID as a temporary tenant-context contract.",
    ):
        _paragraph(document, item, bullet=True)
    _heading(document, "Phase boundary", 2)
    _paragraph(document, "Phase 4A persists project metadata, uploaded file metadata, dataset snapshots, catalogue versions, analysis runs, findings, evidence, exception definitions, and audit events. Phase 4B adds raw row lineage and normalized WBS, budget, cost, commitment, schedule, and progress fact tables.")
    document.save(target)
    return target


def _format_thresholds(rule: dict) -> str:
    values = rule["runtime"].get("thresholds", {})
    materiality = rule["runtime"].get("materiality", {})
    pairs = [f"{key}={value}" for key, value in values.items()]
    pairs.extend(f"{key}={value}" for key, value in materiality.items())
    return "; ".join(pairs) or "No numeric threshold"


def build_catalogue() -> Path:
    target = DOCS / "ControlCheck_AI_Control_Rule_Catalogue_v0.2.docx"
    document = _prepare(
        "ControlCheck_AI_Control_Rule_Catalogue_v0.1.docx",
        target.name,
        "Control Rule Catalogue v0.2",
        "Control Rule Catalogue v0.1",
    )
    catalogue = json.loads(CATALOGUE_PATH.read_text(encoding="utf-8"))
    _heading(document, "Runtime Alignment Supplement v0.2", 1)
    _paragraph(document, "This supplement is generated from data/controlcheck_rule_catalogue_v0.2.json, the governed runtime catalogue consumed by the deterministic engine. It records the evaluation grain and active thresholds for all 20 MVP rules.")
    rows = [
        (
            rule["code"],
            rule["name"],
            rule["runtime"]["evaluation_grain"],
            rule["runtime"]["operator"],
            _format_thresholds(rule),
        )
        for rule in catalogue["rules"]
    ]
    _table(document, ("Code", "Rule", "Grain", "Operator", "Runtime thresholds"), rows)
    _heading(document, "Governed clarifications", 2)
    for item in (
        "CST-004 Vendor Concentration Risk is evaluated at vendor + WBS grain in v0.2, preventing unrelated WBS spend from masking local concentration.",
        "CST-005 High-Value Transaction Outlier requires both inclusive materiality gates: 25% of WBS budget and 3% of project budget.",
        "PRG-003 Cost Rising While Progress Flat requires current-period spend materiality of at least 1% of project budget in addition to its cost/progress movement conditions.",
        "All calculations remain deterministic. LLM output may explain and prioritize findings but cannot change thresholds, metrics, or evidence.",
        "Every persisted result records the catalogue version and SHA-256 so an analysis run can be reproduced against the same governed definition.",
    ):
        _paragraph(document, item, bullet=True)
    document.save(target)
    return target


def build_prd() -> Path:
    target = DOCS / "ControlCheck_AI_PRD_v0.3.docx"
    document = _prepare(
        "ControlCheck_AI_PRD_v0.2.docx",
        target.name,
        "Product Requirements Document v0.3",
        "Product Requirements Document (PRD)",
    )
    _replace(document, "Version 0.2 | MVP Blueprint | 17 August 2026", "Version 0.3 | Backend Persistence Alignment | 17 August 2026")
    _replace_table_row(document, "Version", {1: "0.3"})
    _replace_table_row(document, "Authentication & Organization", {1: "Deferred in Phase 4A", 3: "Tenant isolation now; authentication and complete RBAC later"})
    _replace(document, "16.1 Authentication", "16.1 Authentication (Deferred)")
    _replace(document, "17.1 RBAC", "17.1 RBAC (Deferred)")
    _replace_table_row(document, "AC-08", {1: "Tenant isolation", 2: "Cross-organization project access is rejected server-side; full RBAC is deferred"})
    _replace(
        document,
        "RBAC prevents cross-project/tenant data access.",
        "Phase 4A tenant predicates prevent cross-organization access; full RBAC is deferred.",
    )
    _heading(document, "Phase 4A Product Alignment", 1)
    _paragraph(document, "Phase 4A turns the validated deterministic engine into a durable, organization-scoped backend. It preserves the stateless audit endpoint while adding registered projects, uploaded workbook storage, dataset snapshots, reproducible rule catalogue versions, analysis-run history, persisted findings/evidence, lifecycle status updates, and audit events.")
    _heading(document, "Scope delivered in Phase 4A", 2)
    for item in (
        "PostgreSQL 15+ persistence managed by SQLAlchemy 2 and Alembic migrations.",
        "Local atomic file-storage adapter with path-safety controls and SHA-256 integrity metadata.",
        "Organization-scoped project creation/listing using the temporary X-Organization-ID contract.",
        "Durable analysis execution with Golden and Boundary workbook behavior preserved.",
        "Run history, finding filters, detailed evidence retrieval, and governed finding status transitions.",
        "Stable API error envelope and request ID for operational traceability.",
    ):
        _paragraph(document, item, bullet=True)
    _heading(document, "Explicit deferrals and sequencing", 2)
    _paragraph(document, "Authentication and complete RBAC are deferred to a later hardening phase. The temporary organization header is suitable only for controlled development and validation environments, not public production exposure.")
    _paragraph(document, "Phase 4B will add raw-row lineage and canonical WBS, budget, cost, commitment, schedule, and progress facts. This preserves the Phase 4A delivery focus while preparing database-native analytics and richer exception processing.")
    _paragraph(document, "Frontend implementation remains outside the current phase. The backend contracts should stabilize before UI workflows are built.")
    _heading(document, "Acceptance criteria added in v0.3", 2)
    for item in (
        "Golden workbook creates one succeeded run with 59 findings and source-traceable evidence.",
        "Boundary workbook creates one succeeded run with zero findings.",
        "Workbook/project identity mismatch creates no analysis run.",
        "Engine failure creates one failed run and no findings.",
        "Cross-organization access returns not found and never leaks tenant records.",
        "Alembic upgrade, downgrade, re-upgrade, and zero-drift checks pass against PostgreSQL.",
        "All versioned historical reference artifacts remain byte-for-byte unchanged.",
    ):
        _paragraph(document, item, bullet=True)
    _heading(document, "Change Log", 2)
    _table(
        document,
        ("Version", "Date", "Change"),
        [
            ("0.1", "17 Aug 2026", "Initial MVP blueprint."),
            ("0.2", "17 Aug 2026", "Rule and synthetic validation alignment."),
            ("0.3", "17 Aug 2026", "Phase 4A PostgreSQL persistence, durable API, storage, lifecycle, and Phase 4B sequencing."),
        ],
    )
    document.save(target)
    return target


def main() -> None:
    outputs = (build_erd(), build_catalogue(), build_prd())
    for output in outputs:
        print(output)


if __name__ == "__main__":
    main()
