"""Create ControlCheck PRD v0.2 by minimally editing the preserved v0.1 DOCX."""

from __future__ import annotations

import argparse
from copy import deepcopy
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


def insert_after(anchor: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    element = OxmlElement("w:p")
    anchor._p.addnext(element)
    paragraph = Paragraph(element, anchor._parent)
    if style:
        paragraph.style = style
    if text:
        paragraph.add_run(text)
    return paragraph


def append_blocks(anchor: Paragraph, blocks: list[tuple[str, str]]) -> Paragraph:
    current = anchor
    for style, text in blocks:
        current = insert_after(current, text, style)
    return current


def replace_paragraph_text(paragraph: Paragraph, old: str, new: str) -> None:
    for run in paragraph.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            return
    if old in paragraph.text:
        paragraph.text = paragraph.text.replace(old, new)


def find_paragraph(document, text: str) -> Paragraph:
    for paragraph in document.paragraphs:
        if paragraph.text.strip() == text:
            return paragraph
    raise ValueError(f"Paragraph not found: {text}")


def anchor_before(document, next_heading: str) -> Paragraph:
    paragraphs = document.paragraphs
    for index, paragraph in enumerate(paragraphs):
        if paragraph.text.strip() == next_heading and paragraph.style.name == "Heading 1":
            if index == 0:
                raise ValueError(f"No anchor before {next_heading}")
            return paragraphs[index - 1]
    raise ValueError(f"Heading not found: {next_heading}")


def update_cover_and_metadata(document) -> None:
    replace_paragraph_text(
        document.paragraphs[2],
        "Version 0.1 | MVP Blueprint | 17 August 2026",
        "Version 0.2 | MVP Blueprint | 17 August 2026",
    )
    metadata_table = document.tables[0]
    for row in metadata_table.rows:
        if row.cells[0].text.strip() == "Version":
            row.cells[1].text = "0.2"
            break
    properties = document.core_properties
    properties.title = "ControlCheck AI Product Requirements Document v0.2"
    properties.version = "0.2"
    properties.modified = datetime.now(timezone.utc)


def add_rule_governance(document) -> None:
    anchor = anchor_before(document, "12. Findings & Evidence Engine")
    append_blocks(anchor, [
        ("Heading 2", "11.5 Rule and Threshold Governance"),
        ("Normal", "Every analysis run is bound to an immutable rule-catalogue version. Each structured runtime definition declares its evaluation grain, operator, thresholds, severity bands, lookback, materiality, and permitted exclusions. Dataset, catalogue, and ground-truth versions must be recorded with the run for reproducibility."),
        ("List Bullet", "CST-004 uses the WBS|VENDOR grain so vendor concentration is never ambiguous across work packages."),
        ("List Bullet", "CST-005 triggers only when a transaction is at least 25% of WBS budget and at least 3% of project budget."),
        ("List Bullet", "PRG-003 applies a current-period materiality floor of 1% of project budget before cost-growth logic is evaluated."),
        ("List Bullet", "Threshold or severity changes require a new catalogue version, boundary tests, adjudication rationale, and regression evidence."),
        ("Heading 2", "11.6 Approved Exceptions"),
        ("Normal", "An Approved Exception is an explicit governance record, not silent suppression. It requires an exception ID, rule/entity scope, rationale, approver, effective period, and evidence reference. Expired or unmatched exceptions do not alter findings."),
        ("List Bullet", "Raw results always retain the finding and the applicable exception reference."),
        ("List Bullet", "Exception-aware quality metrics may exclude approved cases from FP/FN counts while raw metrics remain visible."),
        ("List Bullet", "Front-loaded procurement or advance-payment patterns may be approved only through this controlled mechanism."),
    ])


def add_validation_governance(document) -> None:
    anchor = anchor_before(document, "22. Pilot & Go-to-Market")
    append_blocks(anchor, [
        ("Heading 2", "21.3 Ground-Truth Governance"),
        ("Normal", "Ground truth must be curated independently from production rule functions. It is exhaustive for each controlled fixture and stores a unique rule/entity match key, expected severity, metric expectations, evidence anchors, and an adjudication reference. Any label change requires rationale and version history."),
        ("List Bullet", "Golden Positive v0.2 contains 59 exhaustive expected findings across all 20 deterministic rules."),
        ("List Bullet", "Boundary / Negative v0.2 contains 50 literal below, equal, above, and approved-exception cases across 16 numeric rules."),
        ("List Bullet", "Representative or selectively planted labels must not be described as exhaustive ground truth."),
        ("Heading 2", "21.4 Validation Metrics"),
        ("Normal", "Release acceptance requires 100% on controlled validation fixtures for precision, recall, severity accuracy, and metric agreement; zero unreviewed labels; zero unexplained FP/FN; and deterministic repeated output. This is not a customer-accuracy claim and must not be presented as performance on unseen project data."),
        ("List Bullet", "Golden Positive v0.2 verified result: 59 TP, 0 FP, 0 FN, and 100% severity and metric agreement."),
        ("List Bullet", "Boundary / Negative v0.2 verified result: 0 FP, 0 FN, with all 50 manifest cases reconciled."),
        ("List Bullet", "Raw and exception-aware metrics are reported together; approved exceptions remain auditable."),
        ("Heading 2", "21.5 Artifact Version Compatibility"),
        ("Normal", "Before any rule executes, dataset, rule catalogue, and ground truth must share the same major/minor version. A mismatch stops processing, writes no findings output, and returns the stable error code incompatible_artifact_versions through CLI and API."),
    ])


def update_roadmap_and_definition_of_done(document) -> None:
    sprint_heading = find_paragraph(document, "Sprint 3 — Control & Findings")
    append_blocks(sprint_heading, [
        ("List Bullet", "Complete adjudication, Golden Positive, and Boundary / Negative validation alignment before frontend or AI reasoning expansion."),
    ])
    for paragraph in document.paragraphs:
        if paragraph.text.startswith("At least 15") and "golden dataset" in paragraph.text:
            paragraph.text = "All 20 deterministic rules meet the governed v0.2 controlled-fixture acceptance gates before pilot use."
            paragraph.style = "List Bullet"
            break


def add_change_log(document) -> None:
    anchor = find_paragraph(document, "Document Change Log")
    append_blocks(anchor, [
        ("Heading 2", "Change Log v0.2"),
        ("Normal", "Validation Alignment v0.2 formalizes rule and threshold governance, approved exceptions, exhaustive ground-truth governance, controlled-fixture metrics, artifact compatibility, and the validation-before-UI roadmap gate."),
    ])
    table = document.tables[-1]
    template_row = deepcopy(table.rows[-1]._tr)
    table._tbl.append(template_row)
    cells = table.rows[-1].cells
    values = [
        "0.2",
        "17 Aug 2026",
        "Validation Alignment: governed rules, fixtures, ground truth, metrics, exceptions, and version preflight",
        "Product / Founder",
    ]
    for cell, value in zip(cells, values):
        cell.text = value


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("source", type=Path)
    parser.add_argument("output", type=Path)
    args = parser.parse_args()

    document = Document(args.source)
    update_cover_and_metadata(document)
    add_rule_governance(document)
    add_validation_governance(document)
    update_roadmap_and_definition_of_done(document)
    add_change_log(document)
    args.output.parent.mkdir(parents=True, exist_ok=True)
    document.save(args.output)


if __name__ == "__main__":
    main()
